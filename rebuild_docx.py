#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = 'roman_tome1_extracted.txt'
OUT = 'Roman_tome 1.docx'

lines = open(SRC, encoding='utf-8').read().split('\n')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_par(text, align=None, bold=False, size=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    if size:
        r.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

chap_re = re.compile(r'^CHAPITRE (\d+)$')

title_done = False
skip_next = False
for i, line in enumerate(lines):
    if skip_next:
        skip_next = False
        continue
    if not title_done:
        # pages de titre au debut du fichier
        if line == 'LA VIE BELLE':
            add_par(line, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            continue
        if line.startswith('Tome 1'):
            add_par(line, align=WD_ALIGN_PARAGRAPH.CENTER)
            title_done = True
            continue
        if not line.strip():
            continue
        title_done = True
        # premiere ligne de fiction : traiter comme corps de texte

    if line == 'CHAPITRE 1' or chap_re.match(line):
        add_par(line, bold=True, size=24)
        # le sous-titre est la ligne suivante
        if i + 1 < len(lines) and lines[i + 1].strip():
            add_par(lines[i + 1], bold=True, size=24)
            skip_next = True
        continue
    if line == 'ÉPILOGUE':
        add_par(line, bold=True, size=24)
        continue

    add_par(line)

doc.save(OUT)
print('OK:', OUT)
